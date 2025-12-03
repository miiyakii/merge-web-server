from flask import Flask, request, jsonify, send_file, send_from_directory
from flask_cors import CORS
from docx import Document
from docxcompose.composer import Composer
import os
import io

app = Flask(__name__, static_folder='.')

# CORS配置
CORS(app, resources={
    r"/*": {
        "origins": "*",
        "methods": ["GET", "POST", "OPTIONS"],
        "allow_headers": ["Content-Type"]
    }
})

# 模板路径
TEMPLATE_PATH = './证书名字.docx'

# 提供前端页面
@app.route('/')
def index():
    return send_from_directory('.', 'index.html')

@app.route('/generate', methods=['POST'])
def generate_certificates():
    try:
        data = request.json
        names = data.get('names', [])
        
        if not names:
            return jsonify({'error': '没有提供姓名数据'}), 400
        
        if not os.path.exists(TEMPLATE_PATH):
            return jsonify({'error': f'模板文件不存在: {TEMPLATE_PATH}'}), 400
        
        print(f"📝 开始生成证书，共 {len(names)} 人")
        
        first_doc = None
        composer = None
        
        for i, name in enumerate(names):
            doc = Document(TEMPLATE_PATH)
            
            # 替换段落中的 {{name}}
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if "{{name}}" in run.text:
                        run.text = run.text.replace("{{name}}", name)
            
            # 替换表格中的 {{name}}
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if "{{name}}" in run.text:
                                    run.text = run.text.replace("{{name}}", name)
            
            if i == 0:
                first_doc = doc
                composer = Composer(first_doc)
            else:
                composer.append(doc)
            
            print(f"  ✓ 已处理: {name} ({i+1}/{len(names)})")
        
        # 保存到内存
        doc_buffer = io.BytesIO()
        composer.save(doc_buffer)
        doc_buffer.seek(0)
        
        print(f"✅ 生成完成！")
        
        return send_file(
            doc_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='证书合集.docx'
        )
        
    except Exception as e:
        print(f"❌ 错误: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/health', methods=['GET'])
def health_check():
    template_exists = os.path.exists(TEMPLATE_PATH)
    return jsonify({
        'status': 'running',
        'template_exists': template_exists,
        'template_path': TEMPLATE_PATH
    })

if __name__ == '__main__':
    print("=" * 50)
    print("🚀 证书生成服务启动中...")
    print(f"📄 模板路径: {TEMPLATE_PATH}")
    
    if os.path.exists(TEMPLATE_PATH):
        print("✅ 模板文件已找到")
    else:
        print("⚠️  警告: 模板文件不存在，请确保文件位置正确")
    
    print("🌐 服务地址: http://localhost:5001")
    print("📱 Web界面: http://localhost:5001")
    print("=" * 50)
    app.run(debug=True, host='localhost' port=5001)
